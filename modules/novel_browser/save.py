import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt


def _safe_filename(name: str) -> str:
    """Очищає назву файлу від заборонених символів."""
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", name).strip()
    return safe or "chapter"


def _extract_chapter_title(text: str) -> str:
    """
    Автоматично визначає назву глави з перших рядків тексту.
    Для прикладу: 'Розділ 2 - Ліхтар' залишається як є.
    """
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if not lines:
        return "Без назви глави"

    first_line = lines[0]
    
    # 🔹 Якщо перший рядок містить очевидну назву глави - повертаємо його
    # Шукаємо шаблони типу "Розділ X", "Chapter X", "Глава X" тощо
    chapter_patterns = [
        r'^(розділ|chapter|глава|частина)\s+\d+',
        r'^\d+\.',
        r'^\d+$'
    ]
    
    for pattern in chapter_patterns:
        if re.search(pattern, first_line.lower()):
            return first_line

    # 🔹 Якщо є двокрапка — беремо все після неї, але не видаляємо саму назву
    if ":" in first_line:
        parts = first_line.split(":", 1)
        # Повертаємо частину після двокрапки, якщо вона не порожня
        if parts[1].strip():
            return parts[1].strip()

    return first_line or "Без назви глави"


def save_translated_chapter(novel_name: str, chapter_title: str, text: str) -> str:
    """
    Зберігає перекладену главу у форматі .docx у теку 'saved_novels'.
    Назва файлу: <chapter_title>_<дата>.docx
    """
    folder = os.path.join(os.getcwd(), "saved_novels")
    os.makedirs(folder, exist_ok=True)

    # Якщо не передано назву глави — намагаємось знайти її в тексті
    if not chapter_title or chapter_title.strip() in ["", "Без назви глави"]:
        chapter_title = _extract_chapter_title(text)

    safe_chapter = (chapter_title or "Без назви глави").strip()

    # Використовуємо тільки назву глави для файлу
    filename = f"{_safe_filename(safe_chapter)}_{datetime.now().strftime('%Y-%m-%d_%H-%M')}.docx"
    path = os.path.join(folder, filename)

    # Створення документа
    doc = Document()
    
    # Використовуємо повну назву для заголовка в документі
    if novel_name and novel_name.strip():
        full_title = f"{novel_name.strip()} - {safe_chapter}"
    else:
        full_title = safe_chapter
        
    doc.add_heading(full_title, level=1)
    doc.add_paragraph(f"Збережено: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph("")

    # Основний текст
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for p in paragraphs:
        doc.add_paragraph(p)

    doc.save(path)
    return path